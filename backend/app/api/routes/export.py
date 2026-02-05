"""
export.py
API routes for data export functionality
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from typing import Dict, Any, Optional
import structlog
import pandas as pd
import io
import json
from datetime import datetime
import uuid
import base64
from io import BytesIO
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

from app.models.requests import ExportRequest
from app.models.responses import ExportResponse
from app.services.rag_pipeline import rag_pipeline
from app.core.database import db_manager

logger = structlog.get_logger()

router = APIRouter()

# In-memory storage for export data (in production, use Redis or database)
export_cache = {}


@router.post("/export", response_model=ExportResponse)
async def export_data(request: ExportRequest):
    """
    Export data in specified format
    
    Args:
        request: Export request with query, filters, and format
        
    Returns:
        ExportResponse: Export job information
    """
    try:
        export_id = str(uuid.uuid4())
        logger.info("Starting data export", export_id=export_id, format=request.export_format)
        
        # Get data based on request
        if request.query:
            # Use RAG pipeline to get data
            result = await rag_pipeline.process_query(
                user_query=request.query,
                max_results=10000,  # Large limit for export
                language="en"
            )
            sql_results = result.get("retrieved_data", {}).get("sql_results", [])
        elif request.data_filters:
            # Use direct data filters
            sql_results = await _get_data_with_filters(request.data_filters)
        else:
            raise HTTPException(status_code=400, detail="Either query or data_filters must be provided")
        
        if not sql_results:
            raise HTTPException(status_code=404, detail="No data found for export")
        
        # Create export response
        export_response = ExportResponse(
            export_id=export_id,
            format=request.export_format,
            status="completed",
            record_count=len(sql_results),
            created_at=datetime.now(),
            completed_at=datetime.now()
        )
        
        # Store export data for download
        export_data = {
            "export_id": export_id,
            "data": sql_results,
            "metadata": {
                "exported_at": datetime.now().isoformat(),
                "record_count": len(sql_results),
                "format": request.export_format,
                "query": request.query,
                "include_metadata": request.include_metadata
            }
        }
        
        # Store in cache
        export_cache[export_id] = export_data
        
        logger.info("Export completed", export_id=export_id, records=len(sql_results))
        return export_response
        
    except Exception as e:
        logger.error("Export failed", error=str(e), export_id=export_id if 'export_id' in locals() else None)
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.get("/export/{export_id}/download")
async def download_export(export_id: str, format: str = "csv"):
    """
    Download exported data file
    
    Args:
        export_id: Export job ID
        format: File format (csv, json, parquet)
        
    Returns:
        StreamingResponse: File download
    """
    try:
        # Get export data from cache
        export_data = export_cache.get(export_id)
        if not export_data:
            raise HTTPException(status_code=404, detail="Export not found or expired")
        
        data = export_data["data"]
        metadata = export_data["metadata"]
        
        # Convert to DataFrame
        df = pd.DataFrame(data)
        
        # Generate file based on format
        if format.lower() == "csv":
            output = io.StringIO()
            df.to_csv(output, index=False)
            content = output.getvalue()
            media_type = "text/csv"
            filename = f"argo_export_{export_id}.csv"
            
        elif format.lower() == "xlsx":
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='ARGO_Data')
            content = output.getvalue()
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            filename = f"argo_export_{export_id}.xlsx"
            
        elif format.lower() == "json":
            content = json.dumps(data, indent=2, default=str)
            media_type = "application/json"
            filename = f"argo_export_{export_id}.json"
            
        elif format.lower() == "png":
            # Create a visualization and save as PNG
            content = _create_visualization_png(df, metadata)
            media_type = "image/png"
            filename = f"argo_export_{export_id}.png"
            
        elif format.lower() == "docx":
            # Create a Word document with data and visualizations
            content = _create_word_document(df, metadata, data)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"argo_export_{export_id}.docx"
            
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
        
        # Create streaming response
        def generate():
            yield content.encode() if isinstance(content, str) else content
        
        return StreamingResponse(
            generate(),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error("Download failed", error=str(e), export_id=export_id)
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")


async def _get_data_with_filters(filters: Dict[str, Any]) -> list:
    """Get data using direct database filters"""
    try:
        # Build SQL query from filters
        query_parts = ["SELECT * FROM argo_profiles WHERE 1=1"]
        params = []
        
        if filters.get("latitude_min") is not None:
            query_parts.append("AND latitude >= %s")
            params.append(filters["latitude_min"])
        
        if filters.get("latitude_max") is not None:
            query_parts.append("AND latitude <= %s")
            params.append(filters["latitude_max"])
        
        if filters.get("longitude_min") is not None:
            query_parts.append("AND longitude >= %s")
            params.append(filters["longitude_min"])
        
        if filters.get("longitude_max") is not None:
            query_parts.append("AND longitude <= %s")
            params.append(filters["longitude_max"])
        
        if filters.get("date_from"):
            query_parts.append("AND profile_date >= %s")
            params.append(filters["date_from"])
        
        if filters.get("date_to"):
            query_parts.append("AND profile_date <= %s")
            params.append(filters["date_to"])
        
        if filters.get("float_id"):
            query_parts.append("AND float_id = %s")
            params.append(filters["float_id"])
        
        # Add limit
        limit = filters.get("limit", 1000)
        query_parts.append(f"ORDER BY profile_date DESC LIMIT {limit}")
        
        sql_query = " ".join(query_parts)
        
        # Execute query
        results = db_manager.execute_query(sql_query, params)
        return results
        
    except Exception as e:
        logger.error("Failed to get data with filters", error=str(e))
        return []


def _create_visualization_png(df: pd.DataFrame, metadata: Dict[str, Any]) -> bytes:
    """Create a PNG visualization of the data"""
    try:
        # Set up the plot style
        plt.style.use('dark_background')
        fig, axes = plt.subplots(2, 2, figsize=(15, 12))
        fig.suptitle(f'ARGO Ocean Data Analysis\n{metadata.get("query", "Data Export")}', 
                     fontsize=16, color='white')
        
        # Get numeric columns for plotting
        numeric_cols = df.select_dtypes(include=['number']).columns
        
        if len(numeric_cols) >= 1:
            # Temperature plot
            if 'temperature' in str(numeric_cols).lower() or 'temp' in str(numeric_cols).lower():
                temp_cols = [col for col in numeric_cols if 'temp' in col.lower()]
                if temp_cols:
                    df[temp_cols].plot(kind='line', ax=axes[0,0], color='#06b6d4')
                    axes[0,0].set_title('Temperature Data', color='white')
                    axes[0,0].set_ylabel('Temperature (°C)', color='white')
                    axes[0,0].tick_params(colors='white')
            
            # Salinity plot
            if 'salinity' in str(numeric_cols).lower() or 'sal' in str(numeric_cols).lower():
                sal_cols = [col for col in numeric_cols if 'sal' in col.lower()]
                if sal_cols:
                    df[sal_cols].plot(kind='line', ax=axes[0,1], color='#10b981')
                    axes[0,1].set_title('Salinity Data', color='white')
                    axes[0,1].set_ylabel('Salinity (PSU)', color='white')
                    axes[0,1].tick_params(colors='white')
            
            # Data distribution
            if len(numeric_cols) > 0:
                df[numeric_cols[0]].hist(ax=axes[1,0], color='#8b5cf6', alpha=0.7)
                axes[1,0].set_title(f'Distribution of {numeric_cols[0]}', color='white')
                axes[1,0].set_xlabel(numeric_cols[0], color='white')
                axes[1,0].tick_params(colors='white')
            
            # Summary statistics
            axes[1,1].axis('off')
            summary_text = f"""
            Data Summary:
            • Records: {len(df):,}
            • Columns: {len(df.columns)}
            • Numeric columns: {len(numeric_cols)}
            • Export date: {metadata.get('exported_at', 'N/A')}
            • Query: {metadata.get('query', 'N/A')[:50]}...
            """
            axes[1,1].text(0.1, 0.5, summary_text, fontsize=10, color='white',
                          verticalalignment='center', transform=axes[1,1].transAxes)
        
        # Adjust layout and save to bytes
        plt.tight_layout()
        
        # Save to bytes
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight',
                   facecolor='#0f172a', edgecolor='none')
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        plt.close()
        
        return content
        
    except Exception as e:
        logger.error("Failed to create PNG visualization", error=str(e))
        # Return a simple error image
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.text(0.5, 0.5, f'Error creating visualization: {str(e)}', 
                ha='center', va='center', color='white', fontsize=12)
        ax.set_facecolor('#0f172a')
        
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight',
                   facecolor='#0f172a', edgecolor='none')
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        plt.close()
        
        return content


def _create_word_document(df: pd.DataFrame, metadata: Dict[str, Any], data: list) -> bytes:
    """Create a Word document with data and analysis"""
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('ARGO Ocean Data Export Report', 0)
        title.alignment = 1  # Center alignment
        
        # Metadata section
        doc.add_heading('Export Information', level=1)
        doc.add_paragraph(f"Export Date: {metadata.get('exported_at', 'N/A')}")
        doc.add_paragraph(f"Record Count: {metadata.get('record_count', 0):,}")
        doc.add_paragraph(f"Query: {metadata.get('query', 'N/A')}")
        doc.add_paragraph(f"Format: {metadata.get('format', 'N/A').upper()}")
        
        # Data summary
        doc.add_heading('Data Summary', level=1)
        doc.add_paragraph(f"This export contains {len(df):,} records with {len(df.columns)} columns.")
        
        # Column information
        if len(df.columns) > 0:
            doc.add_heading('Column Information', level=2)
            for i, col in enumerate(df.columns[:10]):  # Limit to first 10 columns
                doc.add_paragraph(f"• {col}: {df[col].dtype}")
            
            if len(df.columns) > 10:
                doc.add_paragraph(f"... and {len(df.columns) - 10} more columns")
        
        # Sample data table
        doc.add_heading('Sample Data (First 20 Records)', level=1)
        if len(df) > 0:
            # Create a table with sample data
            sample_df = df.head(20)
            table = doc.add_table(rows=1, cols=min(len(sample_df.columns), 8))
            table.style = 'Table Grid'
            
            # Add headers
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(sample_df.columns[:8]):
                hdr_cells[i].text = str(col)
            
            # Add data rows
            for i, (_, row) in enumerate(sample_df.iterrows()):
                if i >= 20:  # Limit to 20 rows
                    break
                row_cells = table.add_row().cells
                for j, val in enumerate(row[:8]):
                    row_cells[j].text = str(val)[:50]  # Limit cell content length
        
        # Statistics section
        doc.add_heading('Statistical Summary', level=1)
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            stats_df = df[numeric_cols].describe()
            
            # Create statistics table
            stats_table = doc.add_table(rows=1, cols=len(stats_df.columns) + 1)
            stats_table.style = 'Table Grid'
            
            # Add headers
            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = 'Statistic'
            for i, col in enumerate(stats_df.columns):
                hdr_cells[i + 1].text = str(col)
            
            # Add statistics rows
            for stat_name, row in stats_df.iterrows():
                row_cells = stats_table.add_row().cells
                row_cells[0].text = stat_name
                for i, val in enumerate(row):
                    row_cells[i + 1].text = f"{val:.2f}" if pd.notna(val) else "N/A"
        
        # Footer
        doc.add_paragraph("\n" + "="*50)
        doc.add_paragraph("Generated by ARGO FloatChat Export System")
        doc.add_paragraph(f"Export ID: {metadata.get('export_id', 'N/A')}")
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        
        return content
        
    except Exception as e:
        logger.error("Failed to create Word document", error=str(e))
        # Return a simple error document
        doc = Document()
        doc.add_heading('Export Error', 0)
        doc.add_paragraph(f"Failed to create Word document: {str(e)}")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        content = buffer.getvalue()
        buffer.close()
        
        return content
